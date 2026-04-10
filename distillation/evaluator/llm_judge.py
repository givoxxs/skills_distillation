"""LLM Judge: uses Anthropic API to semantically score skill output.

Approach:
  - Extract readable content from the output file (skill-specific)
  - Ask Claude to score 0.0-1.0 against the original prompt + expected_behavior
  - Run N=3 times (default, reduced to 1 during testing via llm_judge_ensemble config)
  - Returns 0.0 on any failure (safe fallback — rule_score still counts)

Usage:
    from evaluator.llm_judge import LLMJudge
    judge = LLMJudge(model="claude-haiku-4-5", ensemble_n=1)
    score, reasoning = judge.score(output_dir, test_case, skill="docx")
"""

from __future__ import annotations

import json
import os
import re
import statistics
from pathlib import Path

import anthropic
from dotenv import load_dotenv

load_dotenv(Path(__file__).parent.parent.parent / ".env")

DEFAULT_MODEL = "claude-haiku-4-5"
DEFAULT_ENSEMBLE = 3
MAX_CONTENT_CHARS = 3000  # truncate extracted content to protect context

JUDGE_SYSTEM = """You are a focused evaluator for AI-generated task outputs.

Score 0.0–1.0:
  1.0 — Output is correct, on-topic, fixture properly handled
  0.7 — Minor issues, mostly correct
  0.4 — Significant content issues or wrong approach
  0.0 — Completely off-topic, hallucinated, or empty

Respond ONLY with JSON:
{"score": <float>, "reasoning": "<2-3 sentences>", "fixture_verdict": "PASS"|"FAIL"|"N/A"}
"""


class LLMJudge:
    def __init__(
        self,
        model: str = DEFAULT_MODEL,
        ensemble_n: int = DEFAULT_ENSEMBLE,
    ) -> None:
        self.model = model
        self.ensemble_n = ensemble_n

    def score(
        self,
        output_dir: str,
        test_case: dict,
        skill: str,
    ) -> tuple[float, str]:
        """Score one test case output.

        Returns:
            (score 0.0-1.0, reasoning string)
        """
        content = _extract_content(output_dir, skill)
        if not content:
            return 0.0, "No extractable content found in output_dir"

        prompt = _build_judge_prompt(test_case, content)

        raw_scores: list[float] = []
        last_reasoning = ""

        for _ in range(self.ensemble_n):
            s, r = _call_claude(prompt, self.model)
            if s >= 0:
                raw_scores.append(s)
                last_reasoning = r

        if not raw_scores:
            return 0.0, "All Claude calls failed"

        final_score = statistics.median(raw_scores)
        return round(final_score, 3), last_reasoning


# ── Content extractors (one per skill) ────────────────────────────────────────


def _extract_content(output_dir: str, skill: str) -> str:
    """Extract human-readable content from output files for LLM Judge."""
    extractors = {
        "docx": _extract_docx,
        "xlsx": _extract_xlsx,
        "slack-gif-creator": _extract_gif_meta,
        "webapp-testing": _extract_script,
        "frontend-design": _extract_html,
        "algorithmic-art": _extract_script,
    }
    fn = extractors.get(skill, _extract_any_text)
    return fn(Path(output_dir))


def _extract_docx(out: Path) -> str:
    files = list(out.rglob("*.docx"))
    if not files:
        return ""
    try:
        from docx import Document

        doc = Document(str(files[0]))
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                style = p.style.name if p.style else "Normal"
                lines.append(f"[{style}] {p.text.strip()}")
        for table in doc.tables:
            for row in table.rows:
                lines.append(
                    "| " + " | ".join(c.text.strip() for c in row.cells) + " |"
                )
        # Extract header/footer content so Judge can verify multi-zone layout
        for i, section in enumerate(doc.sections):
            for zone, label in [(section.header, "Header"), (section.footer, "Footer")]:
                if zone is None:
                    continue
                for p in zone.paragraphs:
                    if p.text.strip():
                        lines.append(f"[{label}-{i}] {p.text.strip()}")
        # Image metadata (Phương án B — format + size, no base64)
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    blob = rel.target_part.blob
                    ctype = rel.target_part.content_type
                    fmt = ctype.split("/")[-1]
                    lines.append(
                        f"[IMAGE] format={fmt}, size={len(blob)} bytes, rId={rel.rId}"
                    )
        except Exception:
            pass
        return _truncate("\n".join(lines))
    except Exception as e:
        return f"(docx parse error: {e})"


def _extract_xlsx(out: Path) -> str:
    files = list(out.rglob("*.xlsx"))
    if not files:
        return ""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(files[0]), data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(max_row=30, values_only=True):
                row_str = " | ".join(str(c) if c is not None else "" for c in row)
                if row_str.strip(" |"):
                    lines.append(row_str)
        return _truncate("\n".join(lines))
    except Exception as e:
        return f"(xlsx parse error: {e})"


def _extract_gif_meta(out: Path) -> str:
    files = list(out.rglob("*.gif"))
    if not files:
        return ""
    lines = []
    for gif_path in files[:3]:
        try:
            from PIL import Image

            img = Image.open(str(gif_path))
            frames = getattr(img, "n_frames", 1)
            size_kb = gif_path.stat().st_size // 1024
            lines.append(
                f"File: {gif_path.name} | Size: {img.size[0]}x{img.size[1]}px | "
                f"Frames: {frames} | File size: {size_kb}KB | Mode: {img.mode}"
            )
        except Exception as e:
            lines.append(f"File: {gif_path.name} | parse error: {e}")
    return "\n".join(lines)


def _extract_html(out: Path) -> str:
    for ext in ("*.html", "*.htm", "*.jsx", "*.tsx"):
        files = list(out.rglob(ext))
        if files:
            return _truncate(files[0].read_text(errors="replace"))
    return _extract_any_text(out)


def _extract_script(out: Path) -> str:
    for ext in ("*.js", "*.ts", "*.py", "*.sh"):
        files = list(out.rglob(ext))
        if files:
            return _truncate(files[0].read_text(errors="replace"))
    return _extract_any_text(out)


def _extract_any_text(out: Path) -> str:
    """Fallback: find any readable text file."""
    for f in sorted(out.rglob("*")):
        if f.is_file() and f.suffix in (".txt", ".md", ".json", ".log"):
            try:
                return _truncate(f.read_text(errors="replace"))
            except Exception:
                continue
    return ""


def _truncate(text: str) -> str:
    if len(text) <= MAX_CONTENT_CHARS:
        return text
    return text[:MAX_CONTENT_CHARS] + f"\n... [truncated, {len(text)} chars total]"


# ── Claude call ───────────────────────────────────────────────────────────────


def _build_judge_prompt(test_case: dict, content: str) -> str:
    cc = test_case.get("content_checks") or {}

    fixture_section = ""
    if test_case.get("fixture_file"):
        fixture_section = f"\nSource fixture: {test_case['fixture_file']}"

    off_topic = f"\n\nVerify content is relevant to this task: {test_case.get('expected_behavior', '')[:200]}"

    fixture_checks = ""
    if cc.get("values_match_fixture"):
        fixture_checks += (
            "\n- Computed values must match actual fixture data (verify numerically)"
        )
    if cc.get("original_text_preserved"):
        fixture_checks += (
            "\n- All original text from the source must be present and intact"
        )

    return f"""Task: {test_case.get('prompt', '')[:500]}
{fixture_section}
{off_topic}
{fixture_checks}

Output content:
---
{content}
---

Score 0.0–1.0. Reply ONLY with JSON: {{"score": <float>, "reasoning": "<2-3 sentences>", "fixture_verdict": "PASS"|"FAIL"|"N/A"}}
"""


def _extract_checklist(test_case: dict) -> list[str]:
    """Build scoring checklist from test_case fields.

    Simplified: keywords are now rules-based in docx_rules.py.
    Only fixture-specific items remain here.
    """
    checklist: list[str] = []
    cc = test_case.get("content_checks") or {}

    if cc.get("values_match_fixture"):
        checklist.append("Computed values match fixture table numerically")
    if cc.get("original_text_preserved"):
        checklist.append("All original text from fixture is preserved")
    if cc.get("output_format") == "json":
        checklist.append("Output is valid JSON")

    # Fallback: expected_behavior (rút gọn)
    expected = test_case.get("expected_behavior", "")
    if expected and len(checklist) < 3:
        lines = re.split(r"[,;]\s*|\n+|(?:\d+\.\s)", expected)
        for line in lines:
            line = line.strip().strip("-•").strip()
            if len(line) > 10:
                checklist.append(line[:120])
                if len(checklist) >= 5:
                    break

    return checklist[:5]


def _call_claude(prompt: str, model: str) -> tuple[float, str]:
    """Call Anthropic API and parse JSON response. Returns (-1, error) on failure."""
    api_key = os.getenv("ANTHROPIC_KEY")
    if not api_key:
        return -1.0, "ANTHROPIC_KEY not set in .env"

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model=model,
            max_tokens=600,
            system=JUDGE_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
        )
        text = message.content[0].text.strip()
    except Exception as e:
        return -1.0, f"API error: {e}"

    return _parse_json_response(text)


def _parse_json_response(text: str) -> tuple[float, str]:
    """Extract score and reasoning from Claude JSON response."""
    clean = re.sub(r"```(?:json)?\s*", "", text).strip().rstrip("`").strip()

    try:
        data = json.loads(clean)
        score = float(data.get("score", -1))
        reasoning = str(data.get("reasoning", ""))

        # Append fixture verdict to reasoning if present
        fixture_verdict = data.get("fixture_verdict", "N/A")
        if fixture_verdict != "N/A":
            reasoning = f"[fixture: {fixture_verdict}] {reasoning}"

        if 0.0 <= score <= 1.0:
            return score, reasoning
        return -1.0, f"Score out of range: {score}"
    except (json.JSONDecodeError, ValueError) as e:
        return -1.0, f"JSON parse error: {e}"
