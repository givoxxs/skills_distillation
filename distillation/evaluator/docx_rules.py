"""Rule-based evaluator for the docx skill (schema v4).

Scoring model
============
Step 1 — Prerequisite gate (must_have_docx):
  If must_have_docx=True (default): file must be a valid, non-empty .docx.
    • file not found → rule_score = 0.0 immediately, skip all checks
    • file corrupt   → rule_score = 0.0 immediately, skip all checks
    • file empty     → rule_score = 0.0 immediately, skip all checks
  If must_have_docx=False: gate skipped (e.g., read → .md or .json).

Step 2 — All checks (prerequisite passed):
  rule_checks + content_checks merged into one flat dict.
  Each field = 1 or more votes (score 0.0 or 1.0).
  rule_score = avg(all votes).

Step 3 — hybrid score:
  hybrid = 0.80 × rule_score + 0.20 × llm_judge_score

LLM Judge (20%) role: off-topic/hallucination detection + fixture integrity.
Run for ALL test cases. ensemble_n default 3 (set 1 in config for cheap testing).

Schema v4 — all fields in one dict
===================================
METADATA:
  id               : str REQUIRED — test identifier, e.g. "tc_a01"
  workflow         : str REQUIRED — create | read | edit | convert
  name             : str REQUIRED — human-readable name
  prompt           : str REQUIRED — exact instruction to the student model
  expected_behavior: str REQUIRED — what correct execution looks like
  skill_gotcha     : str OPTIONAL — critical SKILL.md rule being tested
  fixture_file     : str OPTIONAL — fixture path (read/edit/convert workflows)
  must_have_docx   : bool OPTIONAL — default True. Set False for read-only.

ALL FIELDS BELOW are in ONE MERGED DICT (rule_checks + content_checks):

Format checks (machine-verified):
  xml.contains       : list[str] — substrings that must appear in target XML.
                        Each string = 1 vote.
  xml.absent         : list[str] — substrings that must NOT appear. Each = 1 vote.
  xml.absent_pattern : list[str] — regex patterns (re.DOTALL) must NOT match. Each = 1 vote.
  xml.file           : str — which XML file to check. Default: "word/document.xml".
  xml.footer_contains: list[str] — substrings in word/footer*.xml. Each = 1 vote.
  file.must_exist    : list[str] — paths inside .docx ZIP. Each = 1 vote.
  validate           : bool — run scripts/office/validate.py. 1 vote.
  filename           : str — exact output filename. 1 vote.
  style.header_footer: bool — non-empty header or footer. 1 vote.
  page.min / page.max: int — page count bounds. 1 vote each.
  numbering_references: int — expected distinct w:numId count. 1 vote.

Content checks (rules-based, 1 vote each):
  keywords           : list[str] — must appear in extracted paragraph text.
  keywords_absent    : list[str] — must NOT appear in paragraph text.
  output_format      : str — required file extension, e.g. "json" or "jpg".
  json_keys_from_fixture: bool — JSON keys match fixture table headers.
  output_is_new_file: bool — output filename != fixture filename.

workflow_checks — informational only, NOT in rule_score:
  tool   : str — required tool
  steps  : list[str] — required workflow steps
"""

from __future__ import annotations

import logging
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

from evaluator.base import CheckResult, EvalResult
from evaluator.llm_judge import LLMJudge

_logger = logging.getLogger("distillation.docx_rules")

_VALIDATE_PY = (
    Path(__file__).parent.parent.parent
    / "skill_runner"
    / "skills"
    / "docx"
    / "scripts"
    / "office"
    / "validate.py"
)


class DocxEvaluator:
    """Flat-prerequisite evaluator for .docx output.

    Scoring: prerequisite gate → flat avg of all format + content checks.
    hybrid  = 0.80 × rule_score + 0.20 × llm_judge_score.
    """

    skill = "docx"
    _rule_weight: float = 0.80
    _llm_weight: float = 0.20

    def __init__(
        self,
        judge_model: str = "claude-haiku-4-5",
        use_llm_judge: bool = True,
        ensemble_n: int = 3,
        llm_judge_weight: float = 0.20,
    ) -> None:
        self.use_llm_judge = use_llm_judge
        llm_judge_weight = max(0.0, min(1.0, llm_judge_weight))
        DocxEvaluator._rule_weight = round(1.0 - llm_judge_weight, 6)
        DocxEvaluator._llm_weight = llm_judge_weight
        self._judge = (
            LLMJudge(model=judge_model, ensemble_n=ensemble_n)
            if use_llm_judge
            else None
        )

    # ── Public API ───────────────────────────────────────────────────────────

    def score(
        self,
        output_dir: str,
        test_case: dict,
        model: str,
        round_n: int,
    ) -> EvalResult:
        result = EvalResult(
            test_case_id=test_case["id"],
            skill=self.skill,
            model=model,
            round_n=round_n,
            output_dir=output_dir,
        )

        out = Path(output_dir)
        rc = test_case.get("rule_checks") or {}
        cc = test_case.get("content_checks") or {}
        wc = test_case.get("workflow_checks") or {}
        all_checks_dict = {**rc, **cc}  # merged flat dict

        # ── Locate output file ──────────────────────────────────────────
        fixture_basename = Path(test_case.get("fixture_file", "")).name

        expected_filename = all_checks_dict.get("filename")
        if expected_filename:
            docx_path = out / expected_filename
            if not docx_path.exists():
                # Fallback: pick a docx that isn't the fixture
                candidates = [
                    f for f in out.rglob("*.docx") if f.name != fixture_basename
                ]
                docx_path = (
                    candidates[0]
                    if candidates
                    else next(iter(out.rglob("*.docx")), None)
                )
        else:
            # No expected filename: prefer non-fixture docx over fixture copy
            all_docx = list(out.rglob("*.docx"))
            non_fixture = [f for f in all_docx if f.name != fixture_basename]
            docx_path = (
                non_fixture[0] if non_fixture else (all_docx[0] if all_docx else None)
            )

        # ── Step 1: Prerequisite gate ───────────────────────────────────
        must_have_docx = test_case.get("must_have_docx", True)
        prereq_checks = self._run_prerequisite_checks(docx_path, must_have_docx)
        result.checks = list(prereq_checks)

        if any(c.passed is False for c in prereq_checks):
            result.rule_score = 0.0
        else:
            doc = _try_load(docx_path)
            checks = self._run_all_checks(out, doc, all_checks_dict)
            result.checks += checks
            result.rule_score = _avg(checks)

        # ── workflow_checks (informational, not in rule_score) ───────────
        if wc:
            result.checks += self._run_workflow_checks(out, wc)

        # ── Step 2: hybrid weights (set from config via __init__) ────────
        result._rule_weight = self._rule_weight
        result._llm_weight = self._llm_weight

        # ── LLM Judge (20%) — only if not a total failure ──────────────
        if self._judge and result.rule_score > 0:
            llm_score, reasoning = self._judge.score(output_dir, test_case, self.skill)
            result.llm_judge_score = llm_score
            result.llm_judge_reasoning = reasoning

        return result

    # ── Prerequisite checks ────────────────────────────────────────────────

    def _run_prerequisite_checks(
        self,
        docx_path: Path | None,
        must_have_docx: bool,
    ) -> list[CheckResult]:
        if not must_have_docx:
            _logger.debug("prereq gate: skipped (workflow does not produce .docx)")
            return [
                CheckResult(
                    "must_have_docx",
                    True,
                    1.0,
                    "Prerequisite skipped: workflow does not produce .docx",
                )
            ]

        checks = []
        exists = docx_path is not None and docx_path.exists()
        checks.append(
            CheckResult(
                "file_exists",
                exists,
                1.0 if exists else 0.0,
                "" if exists else "No .docx file found in output_dir",
            )
        )
        _logger.debug("prereq: file_exists=%s path=%s", exists, docx_path)
        if not exists:
            _logger.debug("prereq gate FAILED: file not found — skipping all checks")
            return checks

        doc = _try_load(docx_path)
        parseable = doc is not None
        checks.append(
            CheckResult(
                "file_parseable",
                parseable,
                1.0 if parseable else 0.0,
                ""
                if parseable
                else "python-docx raised an error — file may be corrupt",
            )
        )
        _logger.debug("prereq: file_parseable=%s", parseable)
        if not parseable:
            _logger.debug("prereq gate FAILED: file corrupt — skipping all checks")
            return checks

        size = docx_path.stat().st_size
        non_empty = size > 1024
        checks.append(
            CheckResult(
                "file_not_empty",
                non_empty,
                1.0 if non_empty else 0.0,
                "" if non_empty else f"File only {size} bytes — likely empty",
            )
        )
        _logger.debug("prereq: file_not_empty=%s size=%d bytes", non_empty, size)
        if non_empty:
            _logger.debug("prereq gate PASSED — running all checks")
        return checks

    # ── All checks (format + content merged) ──────────────────────────────

    def _run_all_checks(
        self,
        output_dir: Path,
        doc,
        checks_dict: dict,
    ) -> list[CheckResult]:
        checks: list[CheckResult] = []

        # Baseline: no placeholder text
        if doc is not None:
            ch = self._check_no_placeholders(doc)
            _logger.debug(
                "check %s: passed=%s score=%.1f reason=%r",
                ch.name,
                ch.passed,
                ch.score,
                ch.reason,
            )
            checks.append(ch)

        # style.header_footer
        if checks_dict.get("style.header_footer"):
            ch = self._check_has_header_footer(doc)
            _logger.debug(
                "check %s: passed=%s score=%.1f reason=%r",
                ch.name,
                ch.passed,
                ch.score,
                ch.reason,
            )
            checks.append(ch)

        # Page count
        if checks_dict.get("page.min") or checks_dict.get("page.max"):
            ch = self._check_page_count(
                doc, checks_dict.get("page.min"), checks_dict.get("page.max")
            )
            _logger.debug(
                "check %s: passed=%s score=%.1f reason=%r",
                ch.name,
                ch.passed,
                ch.score,
                ch.reason,
            )
            checks.append(ch)

        # XML / ZIP / validate (requires unzip)
        docx_path = next(iter(output_dir.rglob("*.docx")), None)
        if docx_path and docx_path.exists():
            checks += self._run_xml_checks(docx_path, checks_dict)

        # Content checks (rules-based)
        checks += self._run_content_checks(output_dir, doc, checks_dict)

        _logger.debug(
            "all_checks: %d total, %d passed, rule_score=%.3f",
            len(checks),
            sum(1 for c in checks if c.passed),
            sum(c.score for c in checks) / len(checks) if checks else 0.0,
        )
        return checks

    # ── XML checks ─────────────────────────────────────────────────────────

    def _run_xml_checks(self, docx_path: Path, cd: dict) -> list[CheckResult]:
        checks: list[CheckResult] = []
        need_zip = (
            cd.get("xml.contains")
            or cd.get("xml.absent")
            or cd.get("xml.absent_pattern")
            or cd.get("xml.footer_contains")
            or cd.get("file.must_exist")
            or cd.get("validate")
            or cd.get("numbering_references")
        )

        tmp_dir: Path | None = None
        if need_zip:
            tmp_dir = Path(tempfile.mkdtemp(prefix="docx_eval_"))
            try:
                with zipfile.ZipFile(docx_path, "r") as zf:
                    zf.extractall(tmp_dir)
            except Exception as e:
                checks.append(CheckResult("unzip", False, 0.0, f"Failed to unzip: {e}"))
                return checks

        try:
            xml_file = cd.get("xml.file", "word/document.xml")
            doc_xml = _read_file(tmp_dir / xml_file) if tmp_dir else ""

            # xml.contains — each string = 1 vote
            for needle in cd.get("xml.contains") or []:
                ok = needle in doc_xml
                checks.append(
                    CheckResult(
                        f"xml_has:{needle[:40]}",
                        ok,
                        1.0 if ok else 0.0,
                        "" if ok else f"XML missing in {xml_file}: {needle[:60]}",
                    )
                )

            # xml.absent — each string = 1 vote
            for needle in cd.get("xml.absent") or []:
                ok = needle not in doc_xml
                checks.append(
                    CheckResult(
                        f"xml_absent:{needle[:40]}",
                        ok,
                        1.0 if ok else 0.0,
                        "" if ok else f"XML contains forbidden: {needle[:60]}",
                    )
                )

            # xml.absent_pattern — each regex = 1 vote
            for pattern in cd.get("xml.absent_pattern") or []:
                try:
                    ok = re.compile(pattern, re.DOTALL).search(doc_xml) is None
                except re.error as e:
                    ok = False
                    pattern = f"{pattern} (invalid: {e})"
                checks.append(
                    CheckResult(
                        f"xml_absent_pat:{pattern[:40]}",
                        ok,
                        1.0 if ok else 0.0,
                        "" if ok else f"XML matched forbidden pattern: {pattern[:60]}",
                    )
                )

            # xml.footer_contains
            if cd.get("xml.footer_contains") and tmp_dir:
                footer_xml = "".join(
                    _read_file(f) for f in tmp_dir.glob("word/footer*.xml")
                )
                for needle in cd["xml.footer_contains"]:
                    ok = needle in footer_xml
                    checks.append(
                        CheckResult(
                            f"footer_has:{needle[:40]}",
                            ok,
                            1.0 if ok else 0.0,
                            "" if ok else f"Footer XML missing: {needle[:60]}",
                        )
                    )

            # file.must_exist — each path = 1 vote
            for rel_path in cd.get("file.must_exist") or []:
                if tmp_dir:
                    target = tmp_dir / rel_path.rstrip("/")
                    if rel_path.endswith("/"):
                        ok = target.is_dir() and any(target.iterdir())
                    else:
                        ok = target.exists()
                else:
                    ok = False
                checks.append(
                    CheckResult(
                        f"zip_has:{rel_path}",
                        ok,
                        1.0 if ok else 0.0,
                        "" if ok else f"Missing in docx zip: {rel_path}",
                    )
                )

            # numbering_references — 1 vote
            if cd.get("numbering_references"):
                expected = cd["numbering_references"]
                actual = self._count_numbering_refs(tmp_dir) if tmp_dir else 0
                ok = actual == expected
                checks.append(
                    CheckResult(
                        "numbering_refs",
                        ok,
                        1.0 if ok else 0.0,
                        ""
                        if ok
                        else f"Expected {expected} numbering refs, found {actual}",
                    )
                )

            # validate.py — 1 vote
            if cd.get("validate") and _VALIDATE_PY.exists():
                ch = self._check_validate(docx_path)
                _logger.debug(
                    "check %s: passed=%s score=%.1f reason=%r",
                    ch.name,
                    ch.passed,
                    ch.score,
                    ch.reason,
                )
                checks.append(ch)

        finally:
            if tmp_dir:
                import shutil

                shutil.rmtree(tmp_dir, ignore_errors=True)

        return checks

    # ── Style / structural helpers ─────────────────────────────────────────

    def _check_no_placeholders(self, doc) -> CheckResult:
        if doc is None:
            return CheckResult("no_placeholders", False, 0.0, "Could not load document")
        text = " ".join(p.text for p in doc.paragraphs)
        m = re.compile(
            r"\[INSERT|\[YOUR |\[NAME\]|\[DATE\]|\{\{|TBD\b|TODO\b", re.I
        ).search(text)
        ok = m is None
        return CheckResult(
            "no_placeholders",
            ok,
            1.0 if ok else 0.0,
            "" if ok else f"Unfilled placeholder: '{m.group()}'",
        )

    def _check_has_header_footer(self, doc) -> CheckResult:
        if doc is None:
            return CheckResult(
                "has_header_footer", False, 0.0, "Could not load document"
            )
        has = any(
            (section.header and any(p.text.strip() for p in section.header.paragraphs))
            or (
                section.footer
                and any(p.text.strip() for p in section.footer.paragraphs)
            )
            for section in doc.sections
        )
        return CheckResult(
            "has_header_footer",
            has,
            1.0 if has else 0.0,
            "" if has else "No header or footer with content found",
        )

    def _check_page_count(self, doc, min_pages, max_pages) -> CheckResult:
        if doc is None:
            return CheckResult("page_count", False, 0.0, "Could not load document")
        count = doc.element.body.xml.count('w:type="page"') + 1
        ok = True
        msg = f"Estimated {count} page(s)"
        if min_pages and count < min_pages:
            ok = False
            msg = f"Estimated {count} page(s), expected >= {min_pages}"
        elif max_pages and count > max_pages:
            ok = False
            msg = f"Estimated {count} page(s), expected <= {max_pages}"
        return CheckResult("page_count", ok, 1.0 if ok else 0.0, "" if ok else msg)

    def _count_numbering_refs(self, tmp_dir: Path | None) -> int:
        if tmp_dir is None:
            return 0
        num_path = tmp_dir / "word" / "numbering.xml"
        if not num_path.exists():
            return 0
        try:
            content = num_path.read_text(encoding="utf-8")
            ids = re.findall(r'<w:num\s+w:numId="(\d+)"', content)
            return len(set(ids))
        except Exception:
            return 0

    def _check_validate(self, docx_path: Path) -> CheckResult:
        try:
            proc = subprocess.run(
                [sys.executable, str(_VALIDATE_PY), str(docx_path)],
                capture_output=True,
                text=True,
                timeout=30,
                cwd=str(_VALIDATE_PY.parent),
            )
            ok = proc.returncode == 0
            out = (proc.stdout + proc.stderr).strip()
            msg = "" if ok else out[:800]
            _logger.debug(
                "validate.py: passed=%s rc=%d stdout_len=%d stderr_len=%d",
                ok,
                proc.returncode,
                len(proc.stdout),
                len(proc.stderr),
            )
            if out:
                _logger.debug("validate output: %s", out[:300])
            return CheckResult("validate_passes", ok, 1.0 if ok else 0.0, msg)
        except subprocess.TimeoutExpired:
            _logger.warning("validate.py timed out on %s", docx_path)
            return CheckResult("validate_passes", False, 0.0, "validate.py timed out")
        except Exception as e:
            _logger.warning(
                "validate.py error %s on %s: %s", type(e).__name__, docx_path, e
            )
            return CheckResult("validate_passes", False, 0.0, f"validate.py error: {e}")

    # ── Content checks (rules-based) ──────────────────────────────────────

    def _extract_all_text(self, output_dir: Path, doc, cd: dict) -> str:
        """Extract searchable text from output. Reads doc.paragraphs when available;
        also scans .md/.txt/.json files when doc is None or search_output_files=True."""
        if doc is not None and not cd.get("search_output_files"):
            return " ".join(p.text for p in doc.paragraphs)
        parts: list[str] = []
        if doc is not None:
            parts.append(" ".join(p.text for p in doc.paragraphs))
        for f in sorted(output_dir.rglob("*")):
            if (
                f.is_file()
                and f.suffix in (".md", ".txt", ".json")
                and f.stat().st_size > 0
            ):
                try:
                    parts.append(f.read_text(errors="replace"))
                except Exception:
                    pass
        return " ".join(parts)

    def _run_content_checks(
        self,
        output_dir: Path,
        doc,
        cd: dict,
    ) -> list[CheckResult]:
        checks: list[CheckResult] = []
        all_text = self._extract_all_text(output_dir, doc, cd)
        _logger.debug("content: extracted %d chars from document", len(all_text))

        # keywords — each = 1 vote
        for kw in cd.get("keywords") or []:
            ok = kw in all_text
            ch = CheckResult(
                f"keyword:{kw[:30]}",
                ok,
                1.0 if ok else 0.0,
                "" if ok else f"Keyword not found: '{kw}'",
            )
            _logger.debug(
                "check %s: passed=%s score=%.1f reason=%r",
                ch.name,
                ch.passed,
                ch.score,
                ch.reason,
            )
            checks.append(ch)

        # keywords_absent — each = 1 vote
        for kw in cd.get("keywords_absent") or []:
            ok = kw not in all_text
            ch = CheckResult(
                f"kw_absent:{kw[:30]}",
                ok,
                1.0 if ok else 0.0,
                "" if ok else f"Forbidden keyword found: '{kw}'",
            )
            _logger.debug(
                "check %s: passed=%s score=%.1f reason=%r",
                ch.name,
                ch.passed,
                ch.score,
                ch.reason,
            )
            checks.append(ch)

        # output_format — verify output file type exists
        fmt = cd.get("output_format")
        if fmt:
            ok = bool(list(output_dir.rglob(f"*.{fmt}")))
            ch = CheckResult(
                f"output_format:{fmt}",
                ok,
                1.0 if ok else 0.0,
                "" if ok else f"No .{fmt} file found in output",
            )
            _logger.debug(
                "check %s: passed=%s score=%.1f reason=%r",
                ch.name,
                ch.passed,
                ch.score,
                ch.reason,
            )
            checks.append(ch)

        # json_keys_from_fixture
        if cd.get("json_keys_from_fixture") and cd.get("fixture_file"):
            for sub_ch in self._check_json_keys_from_fixture(
                output_dir, cd["fixture_file"]
            ):
                _logger.debug(
                    "check %s: passed=%s score=%.1f reason=%r",
                    sub_ch.name,
                    sub_ch.passed,
                    sub_ch.score,
                    sub_ch.reason,
                )
                checks.append(sub_ch)

        # output_is_new_file
        if cd.get("output_is_new_file") and cd.get("fixture_file"):
            for sub_ch in self._check_new_file(output_dir, cd["fixture_file"]):
                _logger.debug(
                    "check %s: passed=%s score=%.1f reason=%r",
                    sub_ch.name,
                    sub_ch.passed,
                    sub_ch.score,
                    sub_ch.reason,
                )
                checks.append(sub_ch)

        # values_match_fixture — verify output file exists and is non-empty
        # (numeric verification is delegated to LLM judge; rule check gates on output presence)
        if cd.get("values_match_fixture"):
            ch = self._check_has_output_file(output_dir)
            _logger.debug(
                "check %s: passed=%s score=%.1f reason=%r",
                ch.name,
                ch.passed,
                ch.score,
                ch.reason,
            )
            checks.append(ch)

        # original_text_preserved — sample key phrases from fixture and verify in output
        if cd.get("original_text_preserved") and cd.get("fixture_file"):
            for sub_ch in self._check_original_text_preserved(
                output_dir, doc, cd["fixture_file"]
            ):
                _logger.debug(
                    "check %s: passed=%s score=%.1f reason=%r",
                    sub_ch.name,
                    sub_ch.passed,
                    sub_ch.score,
                    sub_ch.reason,
                )
                checks.append(sub_ch)

        return checks

    def _check_json_keys_from_fixture(
        self, output_dir: Path, fixture_file: str
    ) -> list[CheckResult]:
        fixture_path = self._find_fixture(fixture_file)
        if not fixture_path:
            return [
                CheckResult("json_keys_from_fixture", False, 0.0, "Fixture not found")
            ]
        try:
            from docx import Document

            doc = Document(str(fixture_path))
            header_cells = (
                [c.text.strip() for c in doc.tables[0].rows[0].cells]
                if doc.tables
                else []
            )
            json_files = list(output_dir.rglob("*.json"))
            if not json_files:
                return [
                    CheckResult(
                        "json_keys_from_fixture", False, 0.0, "No JSON output found"
                    )
                ]
            import json as _json

            data = _json.loads(json_files[0].read_text())
            if isinstance(data, list) and data:
                data = data[0]
            json_keys = (
                set(str(k) for k in data.keys()) if isinstance(data, dict) else set()
            )
            fixture_keys = set(header_cells)
            ok = json_keys == fixture_keys
            return [
                CheckResult(
                    "json_keys_from_fixture",
                    ok,
                    1.0 if ok else 0.0,
                    ""
                    if ok
                    else f"JSON keys {json_keys} != fixture headers {fixture_keys}",
                )
            ]
        except Exception as e:
            return [CheckResult("json_keys_from_fixture", False, 0.0, f"Error: {e}")]

    def _check_new_file(self, output_dir: Path, fixture_file: str) -> list[CheckResult]:
        fixture_name = Path(fixture_file).name
        output_files = list(output_dir.rglob("*.docx"))
        ok = any(f.name != fixture_name for f in output_files)
        return [
            CheckResult(
                "output_is_new_file",
                ok,
                1.0 if ok else 0.0,
                "" if ok else f"Output file has same name as fixture: {fixture_name}",
            )
        ]

    def _check_has_output_file(self, output_dir: Path) -> CheckResult:
        """Gate check for values_match_fixture: at least one non-empty output file exists."""
        candidates = list(output_dir.rglob("*"))
        output_files = [
            f
            for f in candidates
            if f.is_file()
            and f.suffix in (".json", ".md", ".txt", ".docx", ".xlsx")
            and f.stat().st_size > 0
        ]
        ok = bool(output_files)
        return CheckResult(
            "has_output_file",
            ok,
            1.0 if ok else 0.0,
            "" if ok else "No output file found — model produced no result",
        )

    def _check_original_text_preserved(
        self, output_dir: Path, doc, fixture_file: str
    ) -> list[CheckResult]:
        """Sample up to 5 unique phrases from fixture and check they appear in output."""
        fixture_path = self._find_fixture(fixture_file)
        if not fixture_path:
            return [
                CheckResult("original_text_preserved", False, 0.0, "Fixture not found")
            ]

        try:
            from docx import Document as _Doc

            fixture_doc = _Doc(str(fixture_path))
            # Collect non-trivial paragraph texts from fixture
            phrases = [
                p.text.strip()
                for p in fixture_doc.paragraphs
                if len(p.text.strip()) > 15
            ][:5]
        except Exception as e:
            return [
                CheckResult(
                    "original_text_preserved", False, 0.0, f"Fixture load error: {e}"
                )
            ]

        if not phrases:
            return [
                CheckResult(
                    "original_text_preserved", True, 1.0, "No phrases to verify"
                )
            ]

        # Get output text — try docx first, fallback to any text file
        output_text = ""
        if doc is not None:
            output_text = " ".join(p.text for p in doc.paragraphs)
        if not output_text:
            for f in sorted(output_dir.rglob("*")):
                if f.is_file() and f.suffix in (".md", ".txt", ".json"):
                    try:
                        output_text = f.read_text(errors="replace")
                        break
                    except Exception:
                        pass

        results = []
        for phrase in phrases:
            ok = phrase in output_text
            results.append(
                CheckResult(
                    f"orig_text:{phrase[:30]}",
                    ok,
                    1.0 if ok else 0.0,
                    "" if ok else f"Fixture phrase missing in output: '{phrase[:60]}'",
                )
            )
        return results

    def _find_fixture(self, fixture_file: str) -> Path | None:
        base = Path(__file__).parent.parent / "test_cases"
        path = base / fixture_file
        return path if path.exists() else None

    # ── workflow_checks (informational only) ───────────────────────────────

    def _run_workflow_checks(self, output_dir: Path, wc: dict) -> list[CheckResult]:
        checks: list[CheckResult] = []
        if wc.get("tool"):
            checks.append(
                CheckResult(
                    f"workflow_tool:{wc['tool']}",
                    True,
                    0.5,
                    f"workflow_checks: tool={wc['tool']} (informational — not in rule_score)",
                )
            )
        if wc.get("steps"):
            checks.append(
                CheckResult(
                    f"workflow_steps:{len(wc['steps'])}",
                    True,
                    0.5,
                    f"workflow_checks: steps={wc['steps']} (informational — not in rule_score)",
                )
            )
        return checks


# ── Module-level helpers ────────────────────────────────────────────────────


def _try_load(path: Path | None):
    if path is None or not path.exists():
        return None
    try:
        from docx import Document

        return Document(str(path))
    except Exception:
        return None


def _read_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


def _avg(checks: list[CheckResult]) -> float:
    if not checks:
        return 0.0
    return sum(c.score for c in checks) / len(checks)
