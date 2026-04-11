#!/usr/bin/env python3
"""Create fixture .docx files for docx skill test cases.

Run from repo root:
    conda activate skills
    python skill_evaluation/test_cases/create_fixtures.py
"""

import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "fixtures"
OUT.mkdir(exist_ok=True)


# ── helpers ───────────────────────────────────────────────────────────────────


def _shade_row(row, hex_color: str = "BDD7EE"):
    """Apply background shading to every cell in a table row."""
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)


# ── fixture 1: simple_report.docx ─────────────────────────────────────────────
# 3 headings (H1/H2/H3), 2 body paragraphs, 1 table 3×3, no images


def make_simple_report():
    doc = Document()

    doc.add_heading("Annual Report 2025", level=1)
    doc.add_heading("Introduction", level=2)
    doc.add_heading("Background Context", level=3)
    doc.add_paragraph(
        "This report summarises the key activities and outcomes of the organisation "
        "during the fiscal year 2025. It covers financial performance, operational "
        "milestones, and strategic initiatives."
    )
    doc.add_paragraph(
        "The data presented here was collected from internal systems and verified "
        "by the finance team. All figures are in USD unless otherwise stated."
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Department", "Budget (USD)", "Spent (USD)"]
    rows_data = [
        ["Engineering", "500,000", "487,000"],
        ["Marketing", "200,000", "198,500"],
        ["Operations", "150,000", "141,200"],
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    path = OUT / "simple_report.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


# ── fixture 2: contract_draft.docx ────────────────────────────────────────────
# sections: Parties, Scope, Payment, Termination; 'Company' x6, 'Party A' x4


def make_contract_draft():
    doc = Document()
    doc.add_heading("Service Agreement", level=1)

    doc.add_heading("1. Parties", level=2)
    doc.add_paragraph(
        "This agreement is entered into between Company (hereinafter 'Company') "
        "and Party A (hereinafter 'Party A'). Company agrees to provide services "
        "as described below."
    )

    doc.add_heading("2. Scope of Work", level=2)
    doc.add_paragraph(
        "Company shall deliver software development services to Party A. "
        "The scope includes design, development, and testing phases. "
        "Company retains ownership of all tooling unless otherwise agreed."
    )

    doc.add_heading("3. Payment Terms", level=2)
    doc.add_paragraph(
        "Party A agrees to pay Company within 30 days of invoice receipt. "
        "Late payments will incur a 1.5% monthly interest charge. "
        "Party A reserves the right to dispute invoices within 10 business days."
    )

    doc.add_heading("4. Termination", level=2)
    doc.add_paragraph(
        "Either party may terminate this agreement with 30 days written notice. "
        "Upon termination, Company shall deliver all completed work to Party A."
    )

    path = OUT / "contract_draft.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


# ── fixture 3: tracked_review.docx ────────────────────────────────────────────
# doc with 3 w:ins tracked insertions by author 'Jane'
# python-docx doesn't support tracked changes, so inject via XML after saving


def make_tracked_review():
    doc = Document()
    doc.add_heading("Project Review Notes", level=1)
    doc.add_paragraph("The project kicked off in January 2025.")
    doc.add_paragraph("Initial milestones were completed on schedule.")
    doc.add_paragraph("The final deliverable was submitted to the client.")

    # Save base doc to temp, then inject tracked changes via XML
    tmp = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(str(tmp))

    # Unzip, patch document.xml, rezip
    tmp_dir = Path(tempfile.mkdtemp())
    with zipfile.ZipFile(tmp, "r") as zf:
        zf.extractall(tmp_dir)

    doc_xml_path = tmp_dir / "word" / "document.xml"
    xml = doc_xml_path.read_text(encoding="utf-8")

    # Inject 3 w:ins elements by appending tracked insertions before </w:body>
    ins_block = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:ins w:id="101" w:author="Jane" w:date="2025-03-01T00:00:00Z">'
        '<w:r><w:t xml:space="preserve"> [Jane: Please clarify this point.]</w:t></w:r>'
        "</w:ins>"
        "</w:p>"
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:ins w:id="102" w:author="Jane" w:date="2025-03-01T00:00:00Z">'
        '<w:r><w:t xml:space="preserve"> [Jane: Add cost breakdown here.]</w:t></w:r>'
        "</w:ins>"
        "</w:p>"
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:ins w:id="103" w:author="Jane" w:date="2025-03-01T00:00:00Z">'
        '<w:r><w:t xml:space="preserve"> [Jane: Reference Appendix B.]</w:t></w:r>'
        "</w:ins>"
        "</w:p>"
    )
    xml = xml.replace("</w:body>", ins_block + "</w:body>")
    doc_xml_path.write_text(xml, encoding="utf-8")

    path = OUT / "tracked_review.docx"
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in tmp_dir.rglob("*"):
            if f.is_file():
                zf.write(f, f.relative_to(tmp_dir))

    shutil.rmtree(tmp_dir)
    tmp.unlink(missing_ok=True)
    print(f"  Created: {path.name}")


# ── fixture 4: newsletter_raw.docx ────────────────────────────────────────────
# plain single-column doc with 3 article blocks, no formatting


def make_newsletter_raw():
    doc = Document()
    doc.add_paragraph("AI Adoption Surges in 2025")
    doc.add_paragraph(
        "Enterprise adoption of AI tools reached an all-time high this quarter, "
        "with over 60% of Fortune 500 companies reporting active deployments."
    )
    doc.add_paragraph(
        "Analysts attribute this growth to falling API costs and improved tooling "
        "for non-technical teams."
    )

    doc.add_paragraph("Remote Work Trends Shift Again")
    doc.add_paragraph(
        "A new survey of 10,000 workers found that hybrid arrangements are now "
        "preferred by 72% of respondents, up from 58% in 2024."
    )
    doc.add_paragraph(
        "Companies offering full flexibility report 23% lower attrition rates "
        "compared to firms requiring five-day office attendance."
    )

    doc.add_paragraph("Open Source Models Close the Gap")
    doc.add_paragraph(
        "Several open-source language models released in Q1 2025 achieved "
        "benchmark scores within 5% of leading proprietary models."
    )
    doc.add_paragraph(
        "This development is expected to accelerate self-hosted deployments "
        "in regulated industries where data residency is required."
    )

    path = OUT / "newsletter_raw.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


# ── fixture 5: data_table.docx ────────────────────────────────────────────────
# single table 5×4 with numeric data, header row shaded


def make_data_table():
    doc = Document()
    doc.add_heading("Quarterly Sales Data", level=1)

    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    headers = ["Product", "Q1 (USD)", "Q2 (USD)", "Q3 (USD)"]
    data = [
        ["Widget A", "45,000", "52,000", "61,000"],
        ["Widget B", "30,000", "28,000", "33,000"],
        ["Widget C", "18,000", "24,000", "27,500"],
        ["Widget D", "12,500", "15,000", "19,000"],
        ["Widget E", "9,000", "11,200", "14,800"],
    ]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    _shade_row(table.rows[0], "BDD7EE")

    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    path = OUT / "data_table.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


# ── main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Creating fixtures in: {OUT}")
    make_simple_report()
    make_contract_draft()
    make_tracked_review()
    make_newsletter_raw()
    make_data_table()
    print("Done.")
